#!/usr/bin/env python3
"""Build a compact Word doc with all before/after edits (uploadable to Google Docs)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parent
SHOTS = ROOT / "shots"
OUT = ROOT / "Inyo-Before-After-Edits.docx"
MAX_W = 6.4


def set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_heading_styled(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1C, 0x1B, 0x19)


def add_body(doc: Document, text: str, *, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x3A, 0x38, 0x34)
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x6B, 0x67, 0x60)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)


def add_image(doc: Document, path: Path, *, width: float | None = None) -> None:
    if not path.exists():
        add_body(doc, f"[Missing image: {path.name}]", italic=True)
        return
    with Image.open(path) as im:
        w_px, h_px = im.size
    target_w = width or MAX_W
    aspect = h_px / max(w_px, 1)
    if aspect < 0.28:
        target_w = MAX_W
    elif aspect > 0.55:
        target_w = min(MAX_W, 5.8)
    doc.add_picture(str(path), width=Inches(target_w))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    last.paragraph_format.space_after = Pt(8)


def main() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_heading_styled(doc, "Inyo — Before & After Edits", 0)
    add_body(
        doc,
        "original/ → exact/  ·  left = before, right = after",
    )

    add_heading_styled(doc, "Edits list", 1)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    headers = ("#", "Edit", "Change")
    rows = [
        ("01", "Favicon", "Black wave + “inyo” text → wave on yellow circle #f5ff66"),
        ("02", "Chat PFPs", "Missing/broken phone header avatar → yellow-circle logo mark"),
        (
            "03",
            "Logo on slide lock",
            "Always opaque → 100% between slides, 0% when a slide locks (home stays visible)",
        ),
        (
            "04",
            "Web Logic scroll-lock",
            "Free scroll → sticky pin; bottom nav hides while locked",
        ),
        (
            "05",
            "Legal titles",
            "{{Title}} — inyo → Privacy Policy / Terms & Conditions — inyo",
        ),
    ]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "F4F1EA")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    sections = [
        {
            "title": "01 — Favicon",
            "body": "Wave mark alone on Inyo’s yellow circle (wordmark removed).",
            "images": [
                ("Before favicon · after favicon · after chat avatar", "compare-favicon-assets.png"),
            ],
        },
        {
            "title": "02 — Chat simulation profile pictures",
            "body": "Phone chat header top-left PFP was missing/broken; replaced with yellow-circle Inyo mark.",
            "images": [
                ("Phone chat header — before | after", "compare-chat-header.png"),
            ],
        },
        {
            "title": "03 — Logo fades on slide lock",
            "body": "Logo stays fully visible while scrolling between slides and on the home slide; fades to 0% when a later slide locks in.",
            "images": [
                ("Logo zoom — before opaque | after hidden", "compare-logo-fade-zoom.png"),
                ("Locked slide viewport — before | after", "compare-slide-lock.png"),
            ],
        },
        {
            "title": "04 — Web Logic scroll-lock",
            "body": "“A simple chat, over a complex web of logic” is now a sticky scroll-lock slide; bottom nav hides while locked; cream edge blend kept.",
            "images": [
                ("Web Logic — before | after", "compare-web-logic.png"),
            ],
        },
        {
            "title": "05 — Legal page titles",
            "body": "Unsubstituted Framer {{Title}} placeholder fixed in exact/ and site/.",
            "images": [
                ("Document <title> — before | after", "compare-privacy-title.png"),
            ],
        },
    ]

    for sec in sections:
        add_heading_styled(doc, sec["title"], 1)
        add_body(doc, sec["body"])
        for caption, fname in sec["images"]:
            add_caption(doc, caption)
            add_image(doc, SHOTS / fname)

    doc.save(OUT)
    print(f"wrote {OUT} ({OUT.stat().st_size / 1024:.0f} KB)")


if __name__ == "__main__":
    main()
